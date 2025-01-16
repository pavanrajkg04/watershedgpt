import os
import glob
import ollama
from langchain.document_loaders import TextLoader, PyPDFLoader, Document
from langchain.embeddings import OpenAIEmbeddings
from langchain.chains import VectorDBQA
from langchain.vectorstores import FAISS
from langchain.chat_models import ChatOpenAI
from langchain.prompts import PromptTemplate


# Step 1: Read documents from a folder (text, PDF, DOCX)
def read_documents_from_folder(folder_path):
    document_texts = []
    file_types = ['*.txt', '*.pdf', '*.docx']  # You can add other file formats

    # Read text files
    for file_type in file_types:
        files = glob.glob(os.path.join(folder_path, file_type))
        for file in files:
            if file.endswith(".txt"):
                with open(file, 'r', encoding='utf-8', errors='ignore') as f:
                    document_texts.append(f.read())
            elif file.endswith(".pdf"):
                with open(file, "rb") as f:
                    loader = PyPDFLoader(f)
                    documents = loader.load()
                    document_texts.extend([doc.page_content for doc in documents])
            elif file.endswith(".docx"):
                from docx import Document
                doc = Document(file)
                for para in doc.paragraphs:
                    document_texts.append(para.text)

    return [Document(page_content=text) for text in document_texts]


# Step 2: Create a vector store using LangChain's FAISS
def create_vector_store(documents):
    embeddings = OpenAIEmbeddings()  # You can change this to other embeddings like Ollama or any LLM-based embeddings
    vector_store = FAISS.from_documents(documents, embeddings)
    return vector_store


# Step 3: Query the chatbot using the vector store
def query_chatbot(user_query, vector_store):
    # Initialize a retriever to fetch relevant documents
    retriever = vector_store.as_retriever(search_type="similarity", search_kwargs={"k": 3})

    # Setup the prompt and the model to answer questions based on the documents
    qa_chain = VectorDBQA.from_chain_type(
        llm=ollama.ChatOpenAI(model="phi4"),  # Replace this with Ollama API
        vectorstore=vector_store,
        chain_type="stuff",  # You can change this to other chain types like "map_reduce" or "refine"
    )

    # Get the response from the model
    response = qa_chain.run(user_query)
    return response


# Step 4: Main function to run the chatbot
def run_chatbot(folder_path):
    # Load documents from the folder
    documents = read_documents_from_folder(folder_path)

    # Create a vector store for efficient querying
    vector_store = create_vector_store(documents)

    print("Bot is ready. Ask me anything!\n")

    while True:
        user_query = input("You: ")
        if user_query.lower() == "exit":
            print("Goodbye!")
            break

        # Query the chatbot
        response = query_chatbot(user_query, vector_store)
        print("Bot:", response)


# Example folder containing your documents
folder_path = 'path/to/your/documents'  # Replace with the path to your documents folder

# Run the chatbot
run_chatbot(folder_path)
